#!/usr/bin/env python3
"""
Reddit Community Monitor - Daily Patrol Report Generator
Uses Reddit RSS feeds (less likely to be blocked) + JSON API for comments.
Falls back to JSON API if RSS fails.
"""

import json
import os
import re
import sys
import time
import datetime
import urllib.request
import urllib.parse
import urllib.error
import ssl
import xml.etree.ElementTree as ET
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from collections import Counter

# For Excel generation
try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
except ImportError:
    os.system(f"{sys.executable} -m pip install openpyxl")
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

# Game subreddits to monitor (with fallbacks)
SUBREDDITS = {
    "Palworld": ["Palworld"],
    "CS2": ["GlobalOffensive"],
    "Valorant": ["Valorant"],
    "LOL": ["leagueoflegends"],
    "DeltaForce": ["DeltaForce"],
    "TFT": ["TeamfightTactics"],
}

# ============================================================
# 爬取需求分类规则（根据用户工作流程图）
# ============================================================
# A类: 工具引流机会帖 - 可以回复并引导用户使用工具
# B类: 用户痛点帖 - 需求收集
TOOL_OPPORTUNITY_KEYWORDS = {
    "DeltaForce": {
        "label": "武器改装/Build相关 → DF工具能解决",
        "keywords": ["weapon build", "gun build", "loadout", "attachment", "build", "改装", "config",
                     "best build", "weapon setup", "gun setup", "ar build", "sniper build", "smg build",
                     "meta build", "optimal build", "weapon combo", "barrel", "stock", "grip", "scope",
                     "muzzle", "foregrip", "laser", "ammo type", "fire rate", "recoil control"],
    },
    "Valorant": {
        "label": "皮肤数据/角色信息 → Val工具能解决",
        "keywords": ["skin", "skins", "cosmetic", "bundle", "weapon skin", "agent info", "agent data",
                     "character info", "skin price", "skin data", "collection", "battle pass",
                     "agent ability", "agent stats", "agent guide", "星辉", "皮肤", "角色数据",
                     "valuant skin", "skin list", "cosmetics list", "agent list", "contract"],
    },
    "Palworld": {
        "label": "配种组合/繁殖 → Palworld工具能解决",
        "keywords": ["breeding", "breed", "breed combo", "pairing", "pal combination", "breeding chain",
                     "breeding recipe", "breed guide", "how to breed", "breed result", "offspring",
                     "配种", "繁殖", "breeding calculator", "breeding tree", "pal breeding",
                     "child pal", "parent pal", "inherit", "passive skill", "iv", "stats",
                     "alpha", "lucky", "fusion", "perfect pal", "soul"],
    },
    "LOL": {
        "label": "ARAM攻略/阵容 → LOL工具能解决",
        "keywords": ["aram", "aram guide", "aram comp", "aram build", "aram tier", "aram strategy",
                     "aram best", "aram champion", "all random", "aram tier list", "aram pick",
                     "aram counter", "aram comp guide", "aram win", "random build", "ARAM",
                     "aram meta", "aram op", "aram guide"],
    },
    "TFT": {
        "label": "TFT阵容/Tier List → TFT工具能解决",
        "keywords": ["tier list", "tier", "comp", "composition", "best comp", "meta comp",
                     "team comp", "build guide", "tier", "S tier", "A tier", "tier list",
                     "meta", "patch notes", "best build", "op comp", "meta build",
                     "comp guide", "re Roll", "fast", "positioning", "augment tier",
                     "teamfight tactics", "set", "comp tier", "tier comp"],
    },
    "CS2": {
        "label": "CS2数据/统计 → CS2工具能解决",
        "keywords": ["stats", "statistics", "data", "kd", "k/d", "winrate", "win rate",
                     "match data", "player stats", "weapon stats", "skin price", "skin value",
                     "inventory value", "float", "wear", "pattern", "skin data",
                     "market price", "stat track", "competitive stats", "premier stats",
                     "rank distribution", "elo", "rating", "faceit level"],
    },
}

# B类: 用户痛点关键词
PAIN_POINT_KEYWORDS = {
    "all": {
        "complaint": ["broken", "bug", "bugged", "glitch", "crash", "lag", "unplayable",
                      "worst", "terrible", "awful", "frustrating", "annoying", "hate",
                      "fix this", "please fix", "needs to fix", "so bad", "ruined",
                      "disappointed", "unfair", "rip off", "scam", "pay to win", "p2w"],
        "feature_request": ["wish there was", "need a tool", "is there a way", "is there a tool",
                            "can someone help", "how do i", "is it possible", "feature request",
                            "would be nice if", "i wish", "why is there no", "missing feature",
                            "why can't i", "why doesn't", "any way to", "is there any"],
        "data_error": ["wrong data", "incorrect", "inaccurate", "outdated", "not updating",
                       "data error", "stale data", "broken data", "missing data"],
        "unmet_need": ["no solution", "nobody knows", "can't find", "no guide",
                       "no tool for", "nothing helps", "still stuck", "giving up",
                       "no one answered", "unresolved", "still looking"],
    }
}

# ============================================================
# Facebook 搜索配置
# ============================================================
FB_SEARCH_QUERIES = {
    "Palworld": "Palworld",
    "CS2": "Counter-Strike 2 CS2",
    "Valorant": "Valorant",
    "LOL": "League of Legends",
    "DeltaForce": "Delta Force game",
    "TFT": "Teamfight Tactics TFT",
}

# ============================================================
# X (Twitter) 搜索配置
# ============================================================
X_SEARCH_QUERIES = {
    "Palworld": "Palworld",
    "CS2": "CS2 Counter-Strike",
    "Valorant": "Valorant",
    "LOL": "League of Legends LOL",
    "DeltaForce": "Delta Force game",
    "TFT": "Teamfight Tactics TFT",
}


def _get_search_engines(site, query):
    """Return a list of search URLs from different engines to try, to avoid rate limiting."""
    q = query.replace(' ', '+')
    # Multiple search engines + DuckDuckGo with different endpoints
    return [
        # DuckDuckGo Lite (different endpoint)
        f"https://lite.duckduckgo.com/lite/?q=site:{site}+{q}",
        # DuckDuckGo HTML
        f"https://html.duckduckgo.com/html/?q=site:{site}+{q}",
        # Startpage (proxied Google results)
        f"https://www.startpage.com/sp/search?q=site:{site}+{q}",
        # Ecosia
        f"https://www.ecosia.org/search?q=site:{site}+{q}",
    ]


def _get_random_ua():
    """Return a random User-Agent to avoid fingerprinting."""
    uas = [
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36",
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:128.0) Gecko/20100101 Firefox/128.0",
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 14_5) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.5 Safari/605.1.15",
        "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/130.0.0.0 Safari/537.36",
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/128.0.0.0 Safari/537.36",
    ]
    import random
    return random.choice(uas)


def _search_multi_engine(site, query, limit=15, search_type="FB"):
    """Search using multiple engines with UA rotation and delays to avoid rate limiting."""
    import re as _re
    import random
    
    search_urls = _get_search_engines(site, query)
    
    for attempt, search_url in enumerate(search_urls):
        # Random delay between 8-15 seconds to avoid rate limiting
        delay = random.uniform(8, 15)
        if attempt == 0:
            delay = 3  # First attempt faster
        print(f"  {search_type} Search (attempt {attempt+1}/{len(search_urls)}, delay {delay:.0f}s): {search_url[:80]}...")
        time.sleep(delay)
        
        headers = {"User-Agent": _get_random_ua(), "Accept": "text/html"}
        status, content = fetch_url(search_url, retries=2, delay=5, custom_headers=headers)
        if not content:
            print(f"  {search_type} attempt {attempt+1} failed (no content)")
            continue
        
        text = content.decode("utf-8", errors="replace")
        
        # Parse results - try different patterns for different search engines
        escaped_site = site.replace('.', r'\.')
        
        # DuckDuckGo Lite format
        links = _re.findall(r'class="result-link"[^>]*href="([^"]*' + escaped_site + r'[^"]*)"', text)
        titles = _re.findall(r'class="result-snippet"[^>]*>(.*?)</td>', text, _re.DOTALL)
        
        # DuckDuckGo HTML format
        if not links:
            links = _re.findall(r'class="result__a"[^>]*href="([^"]*' + escaped_site + r'[^"]*)"', text)
            titles = _re.findall(r'class="result__a"[^>]*>(.*?)</a>', text, _re.DOTALL)
        snippets = _re.findall(r'class="result__snippet"[^>]*>(.*?)</[^>]+>', text, _re.DOTALL)
        
        # Startpage format
        if not links:
            links = _re.findall(r'href="([^"]*' + escaped_site + r'[^"]*)"[^>]*class="result-title"', text)
            titles = _re.findall(r'class="result-title"[^>]*>(.*?)</a>', text, _re.DOTALL)
            snippets = _re.findall(r'class="result-excerpt"[^>]*>(.*?)</p>', text, _re.DOTALL)
        
        # Ecosia format
        if not links:
            links = _re.findall(r'href="([^"]*' + escaped_site + r'[^"]*)"[^>]*', text)
            titles = _re.findall(r'<span[^>]*>(.*?)</span>', text, _re.DOTALL)
            snippets = _re.findall(r'class="snippet"[^>]*>(.*?)</[^>]+>', text, _re.DOTALL)
        
        if links:
            print(f"  {search_type}: Found {len(links)} links on attempt {attempt+1}")
            return links, titles, snippets
    
    print(f"  {search_type}: All {len(search_urls)} engines exhausted, no results")
    return [], [], []


def search_facebook_posts(game, query, limit=15):
    """Search for Facebook posts about a game using multiple search engines."""
    links, titles, snippets = _search_multi_engine("facebook.com", query, limit, "FB")
    if not links:
        print(f"  FB: Found 0 results for {game}")
        return []
    
    import re as _re
    import urllib.parse
    posts = []
    for i in range(min(len(links), limit)):
        url = links[i].replace("&amp;", "&")
        if "uddg=" in url:
            parsed = urllib.parse.parse_qs(urllib.parse.urlparse(url).query)
            if "uddg" in parsed:
                url = parsed["uddg"][0]
        
        title = _re.sub(r'<[^>]+>', '', titles[i]).strip() if i < len(titles) else ""
        snippet = _re.sub(r'<[^>]+>', '', snippets[i]).strip() if i < len(snippets) else ""
        
        post = {
            "title": title[:200], "url": url, "author": "Facebook User",
            "score": 0, "num_comments": 0, "upvote_ratio": 0,
            "selftext": snippet[:500], "link_flair_text": "Facebook",
            "subreddit": "facebook", "sort_type": "facebook", "created_utc": 0,
        }
        classify_post(post, game)
        posts.append(post)
    
    print(f"  FB: Found {len(posts)} results for {game}")
    return posts


def search_x_posts(game, query, limit=15):
    """Search for X (Twitter) posts using multiple search engines."""
    # Try x.com first, then twitter.com as fallback
    links, titles, snippets = _search_multi_engine("x.com", query, limit, "X")
    if not links:
        print(f"  X: Trying twitter.com fallback for {game}")
        links, titles, snippets = _search_multi_engine("twitter.com", query, limit, "X")
    
    if not links:
        print(f"  X: Found 0 results for {game}")
        return []
    
    import re as _re
    import urllib.parse
    posts = []
    for i in range(min(len(links), limit)):
        url = links[i].replace("&amp;", "&")
        if "uddg=" in url:
            parsed = urllib.parse.parse_qs(urllib.parse.urlparse(url).query)
            if "uddg" in parsed:
                url = parsed["uddg"][0]
        
        title = _re.sub(r'<[^>]+>', '', titles[i]).strip() if i < len(titles) else ""
        snippet = _re.sub(r'<[^>]+>', '', snippets[i]).strip() if i < len(snippets) else ""
        
        post = {
            "title": title[:200], "url": url, "author": "X User",
            "score": 0, "num_comments": 0, "upvote_ratio": 0,
            "selftext": snippet[:500], "link_flair_text": "X/Twitter",
            "subreddit": "x_twitter", "sort_type": "x", "created_utc": 0,
        }
        classify_post(post, game)
        posts.append(post)
    
    print(f"  X: Found {len(posts)} results for {game}")
    return posts


def fetch_x_via_syndication(game, screen_names=None, limit=15):
    """Fetch X/Twitter posts via syndication API (no auth, no IP blocking)."""
    import json as _json
    posts = []
    
    x_accounts = {
        "Palworld": ["Palworld_EN", "PalworldGame"],
        "CS2": ["csgo_dev", "CS2"],
        "Valorant": ["PlayVALORANT", "valorantsource"],
        "LOL": ["LeagueOfLegends", "lolesports"],
        "DeltaForce": ["DeltaForceGame", "TencentGames"],
        "TFT": ["TFT", "TFTesports"],
    }
    
    accounts = screen_names or x_accounts.get(game, [game])
    
    for account in accounts:
        url = f"https://cdn.syndication.twimg.com/timeline/profile?screen_name={account}&count=10"
        print(f"  X API: @{account}")
        status, content = fetch_url(url, retries=2, delay=3)
        if not content:
            print(f"  X API failed for @{account}")
            continue
        
        try:
            data = _json.loads(content.decode("utf-8", errors="replace"))
            tweets = data.get("tweets", [])
            for tweet in tweets[:limit]:
                text = tweet.get("full_text", tweet.get("text", ""))
                tweet_id = tweet.get("id_str", tweet.get("id", ""))
                
                post = {
                    "title": text[:200],
                    "url": f"https://x.com/{account}/status/{tweet_id}",
                    "author": f"@{account}",
                    "score": 0,
                    "num_comments": 0,
                    "upvote_ratio": 0,
                    "selftext": text[:500],
                    "link_flair_text": "X/Twitter",
                    "subreddit": "x_twitter",
                    "sort_type": "x",
                    "created_utc": 0,
                }
                classify_post(post, game)
                posts.append(post)
        except Exception as e:
            print(f"  X API parse error for @{account}: {e}")
    
    print(f"  X (API): Found {len(posts)} posts for {game}")
    return posts


def fetch_fb_x_via_reddit(game, query, limit=10):
    """Find FB/X content by searching Reddit for cross-posts linking to FB/X."""
    import json as _json
    search_terms = [f"url:facebook.com {query}", f"url:x.com {query}", f"url:twitter.com {query}"]
    posts = []
    
    for term in search_terms:
        url = f"https://www.reddit.com/search.json?q={urllib.parse.quote(term)}&limit=5&sort=new"
        status, content = fetch_url(url, retries=2, delay=3)
        if not content:
            continue
        
        try:
            data = _json.loads(content.decode("utf-8", errors="replace"))
            children = data.get("data", {}).get("children", [])
            for child in children:
                rd = child.get("data", {})
                post_url = rd.get("url", "")
                if "facebook.com" in post_url or "x.com" in post_url or "twitter.com" in post_url:
                    is_fb = "facebook.com" in post_url
                    post = {
                        "title": rd.get("title", "")[:200],
                        "url": post_url,
                        "author": rd.get("author", ""),
                        "score": rd.get("score", 0),
                        "num_comments": rd.get("num_comments", 0),
                        "upvote_ratio": rd.get("upvote_ratio", 0),
                        "selftext": rd.get("selftext", "")[:500],
                        "link_flair_text": "Facebook" if is_fb else "X/Twitter",
                        "subreddit": rd.get("subreddit", ""),
                        "sort_type": "facebook" if is_fb else "x",
                        "created_utc": rd.get("created_utc", 0),
                    }
                    classify_post(post, game)
                    posts.append(post)
        except Exception as e:
            print(f"  Reddit crosspost error: {e}")
    
    print(f"  Reddit crosspost: Found {len(posts)} FB/X posts for {game}")
    return posts


def fetch_fb_x_via_playwright(game, fb_query, x_query, limit=10):
    """Fetch FB/X posts using Playwright browser (bypasses search engine blocking)."""
    posts = []
    
    try:
        from playwright.sync_api import sync_playwright
        import time as _time
        import random
        
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True, args=["--no-sandbox", "--disable-gpu"])
            context = browser.new_context(
                user_agent=random.choice([
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36",
                    "Mozilla/5.0 (Macintosh; Intel Mac OS X 14_5) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.5 Safari/605.1.15",
                    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/130.0.0.0 Safari/537.36",
                ]),
                viewport={"width": 1920, "height": 1080},
                locale="en-US",
            )
            
            # === Facebook scraping ===
            fb_urls_to_try = [
                f"https://www.facebook.com/search/posts?q={fb_query.replace(' ', '+')}",
                f"https://www.facebook.com/hashtag/{fb_query.split()[0].lower()}",
            ]
            
            # Try Facebook public group pages
            fb_group_map = {
                "Palworld": "https://www.facebook.com/groups/palworld",
                "CS2": "https://www.facebook.com/groups/counterstrike",
                "Valorant": "https://www.facebook.com/PlayVALORANT",
                "LOL": "https://www.facebook.com/leagueoflegends",
                "DeltaForce": "https://www.facebook.com/DeltaForceGame",
                "TFT": "https://www.facebook.com/TeamfightTactics",
            }
            
            fb_url = fb_group_map.get(game, fb_urls_to_try[0])
            print(f"  FB Playwright: {fb_url[:60]}...")
            
            page = context.new_page()
            page.set_default_timeout(15000)
            
            try:
                page.goto(fb_url, wait_until="domcontentloaded")
                _time.sleep(random.uniform(3, 5))
                
                # Scroll to load more content
                for _ in range(3):
                    page.evaluate("window.scrollBy(0, 1000)")
                    _time.sleep(1)
                
                # Extract posts from FB page
                content = page.content()
                import re as _re
                
                # FB posts are in div elements with data attributes
                # Try multiple selectors
                post_elements = page.query_selector_all('[data-ad-comet-preview="message"]')
                if not post_elements:
                    post_elements = page.query_selector_all('[dir="auto"] > span')
                if not post_elements:
                    post_elements = page.query_selector_all('div[data-testid="post_message"]')
                
                seen = set()
                for elem in post_elements[:limit*2]:
                    try:
                        text = elem.inner_text().strip()
                        if len(text) > 20 and text not in seen:
                            seen.add(text)
                            post = {
                                "title": text[:200],
                                "url": fb_url,
                                "author": "Facebook User",
                                "score": 0,
                                "num_comments": 0,
                                "upvote_ratio": 0,
                                "selftext": text[:500],
                                "link_flair_text": "Facebook",
                                "subreddit": "facebook",
                                "sort_type": "facebook",
                                "created_utc": 0,
                            }
                            classify_post(post, game)
                            posts.append(post)
                            if len([p for p in posts if p["sort_type"] == "facebook"]) >= limit:
                                break
                    except:
                        continue
                
                fb_count = len([p for p in posts if p["sort_type"] == "facebook"])
                print(f"  FB Playwright: Found {fb_count} posts")
            except Exception as e:
                print(f"  FB Playwright error: {e}")
            finally:
                page.close()
            
            # === X/Twitter scraping ===
            x_account_map = {
                "Palworld": "Palworld_EN",
                "CS2": "csgo_dev",
                "Valorant": "PlayVALORANT",
                "LOL": "LeagueOfLegends",
                "DeltaForce": "DeltaForceGame",
                "TFT": "TFT",
            }
            
            x_account = x_account_map.get(game, game)
            x_url = f"https://x.com/{x_account}"
            print(f"  X Playwright: {x_url}")
            
            page2 = context.new_page()
            page2.set_default_timeout(15000)
            
            try:
                page2.goto(x_url, wait_until="domcontentloaded")
                _time.sleep(random.uniform(3, 5))
                
                # Scroll to load tweets
                for _ in range(3):
                    page2.evaluate("window.scrollBy(0, 1000)")
                    _time.sleep(1)
                
                # Extract tweets
                tweet_elements = page2.query_selector_all('[data-testid="tweetText"]')
                
                seen = set()
                for elem in tweet_elements[:limit*2]:
                    try:
                        text = elem.inner_text().strip()
                        if len(text) > 10 and text not in seen:
                            seen.add(text)
                            # Try to find the tweet URL from parent element
                            tweet_url = x_url
                            try:
                                link = elem.evaluate("el => { let parent = el.parentElement; while(parent) { let a = parent.querySelector('a[href*="/status/"]'); if(a) return a.href; parent = parent.parentElement; } return ''; }")
                                if link:
                                    tweet_url = link
                            except:
                                pass
                            
                            post = {
                                "title": text[:200],
                                "url": tweet_url,
                                "author": f"@{x_account}",
                                "score": 0,
                                "num_comments": 0,
                                "upvote_ratio": 0,
                                "selftext": text[:500],
                                "link_flair_text": "X/Twitter",
                                "subreddit": "x_twitter",
                                "sort_type": "x",
                                "created_utc": 0,
                            }
                            classify_post(post, game)
                            posts.append(post)
                            if len([p for p in posts if p["sort_type"] == "x"]) >= limit:
                                break
                    except:
                        continue
                
                x_count = len([p for p in posts if p["sort_type"] == "x"])
                print(f"  X Playwright: Found {x_count} posts")
            except Exception as e:
                print(f"  X Playwright error: {e}")
            finally:
                page2.close()
            
            browser.close()
            
    except ImportError:
        print("  Playwright not installed, skipping browser scraping")
    except Exception as e:
        print(f"  Playwright error: {e}")
    
    total_fb = len([p for p in posts if p["sort_type"] == "facebook"])
    total_x = len([p for p in posts if p["sort_type"] == "x"])
    print(f"  Playwright total: FB={total_fb} X={total_x} for {game}")
    return posts



POSTS_PER_SUB = 25
COMMENTS_PER_POST = 10

ctx = ssl.create_default_context()
ctx.check_hostname = False
ctx.verify_mode = ssl.CERT_NONE

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36",
    "Accept": "*/*",
    "Accept-Language": "en-US,en;q=0.5",
}


def fetch_url(url, retries=3, delay=5, custom_headers=None):
    """Fetch URL content with retries. Returns (status_code, content_bytes)."""
    for attempt in range(retries):
        try:
            req_headers = custom_headers if custom_headers else HEADERS
            req = urllib.request.Request(url, headers=req_headers)
            with urllib.request.urlopen(req, context=ctx, timeout=30) as resp:
                return resp.status, resp.read()
        except urllib.error.HTTPError as e:
            print(f"  HTTP {e.code}: {e.reason} (attempt {attempt+1}/{retries})", file=sys.stderr)
            if e.code in (429, 503):
                wait = delay * (attempt + 1) * 2
                print(f"  Rate limited, waiting {wait}s...", file=sys.stderr)
                time.sleep(wait)
            elif attempt < retries - 1:
                time.sleep(delay)
        except Exception as e:
            print(f"  Error: {e} (attempt {attempt+1}/{retries})", file=sys.stderr)
            if attempt < retries - 1:
                time.sleep(delay)
    return None, None


def fetch_json(url, retries=3, delay=5):
    """Fetch JSON from URL with retries."""
    status, content = fetch_url(url, retries, delay)
    if content:
        try:
            return json.loads(content.decode("utf-8"))
        except json.JSONDecodeError as e:
            print(f"  JSON parse error: {e}", file=sys.stderr)
    return None


# ============================================================
# 主贴内容总结配置
# ============================================================
# Post type detection keywords
POST_TYPE_KEYWORDS = {
    "新闻资讯": ["update", "patch", "news", "announce", "release", "launch", "season", "event", "maintenance", "downtime", "server", "official"],
    "攻略教程": ["guide", "how to", "tips", "trick", "best", "build", "tier list", "meta", "strategy", "walkthrough", "tutorial", "explain"],
    "问题求助": ["help", "how do", "why", "question", "stuck", "can't", "cannot", "issue", "problem", "fix", "bug", "crash", "error"],
    "讨论交流": ["discuss", "thought", "opinion", "what do you think", "anyone else", "does anyone", "am i the only", "hot take", "unpopular opinion"],
    "内容分享": ["look at", "check out", "my", "i made", "i drew", "fan art", "clip", "video", "screenshot", "showcase", "share"],
    "赛事竞技": ["tournament", "esports", "vct", "major", "championship", "pro", "team", "match", "vs", "vs.", "score", "bracket", "qualifier"],
    "反馈建议": ["suggestion", "feedback", "wish", "should add", "need", "would be nice", "please add", "feature request", "idea"],
    "吐槽抱怨": ["broken", "terrible", "awful", "worst", "hate", "annoying", "frustrating", "unfair", "ridiculous", "stupid", "trash"],
}

# Topic detection keywords by category
TOPIC_KEYWORDS = {
    "武器装备": ["weapon", "gun", "rifle", "skin", "knife", "loadout", "attachment", "crosshair", "accuracy", "damage", "fire rate"],
    "角色英雄": ["agent", "character", "hero", "champion", "ability", "ultimate", "skill", "passive", "rework", "buff", "nerf"],
    "竞技对战": ["rank", "ranked", "matchmaking", "mmr", "elo", "diamond", "immortal", "radiant", "gold", "silver", "platinum", "master"],
    "游戏更新": ["patch", "update", "version", "changelog", "hotfix", "maintenance", "season", "chapter", "battle pass"],
    "经济系统": ["price", "market", "economy", "cost", "buy", "sell", "trade", "currency", "gold", "coin", "store", "shop"],
    "地图场景": ["map", "location", "zone", "area", "spot", "position", "site", "territory", "poi", "point of interest"],
    "社交社区": ["community", "friend", "guild", "clan", "team", "party", "squad", "multiplayer", "co-op", "coop", "online"],
    "画质性能": ["fps", "lag", "stutter", "crash", "bug", "glitch", "graphics", "settings", "optimization", "performance", "driver"],
    "角色培养": ["breeding", "level", "upgrade", "evolve", "skill tree", "talent", "mastery", "progression", "build"],
    "赛事电竞": ["tournament", "esports", "pro", "vct", "major", "championship", "lec", "lcs", "lck", "vcs", "worlds", "international"],
    "MOD创意": ["mod", "custom", "creative", "sandbox", "build", "design", "art", "creation", "blueprint"],
    "新手相关": ["newbie", "beginner", "first time", "starter", "new player", "guide", "how to start"],
    "故事剧情": ["lore", "story", "quest", "campaign", "narrative", "cutscene", "ending", "character background"],
    "皮肤外观": ["skin", "cosmetic", "outfit", "emote", "spray", "card", "banner", "appearance", "customize"],
    "配装策略": ["comp", "composition", "meta", "tier list", "synergy", "team comp", "build order", "strategy"],
    "操作技巧": ["mechanics", "aim", "movement", "utility", "smoke", "flashbang", "grenade", "recoil", "spray", "flick"],
}


def summarize_post_content(post, game):
    """Generate a detailed Chinese summary of what a post is discussing.
    
    Analyzes title, selftext, and comments to determine:
    - Post type (news, guide, question, discussion, etc.)
    - Main topic (weapons, agents, competitive, updates, etc.)
    - Key discussion points from comments
    - Sentiment and engagement level
    """
    title = post.get("title", "")
    selftext = post.get("selftext", "")
    comments = post.get("comments", [])
    flair = post.get("link_flair_text", "")
    score = post.get("score", 0)
    num_comments = post.get("num_comments", 0)
    
    combined = f"{title} {selftext}".lower()
    
    # 1. Detect post type
    post_type = "综合讨论"
    for ptype, keywords in POST_TYPE_KEYWORDS.items():
        matches = [kw for kw in keywords if kw in combined]
        if matches:
            post_type = ptype
            break
    
    # Use flair as hint if available
    if flair:
        flair_lower = flair.lower()
        if any(k in flair_lower for k in ["news", "update", "announcement"]):
            post_type = "新闻资讯"
        elif any(k in flair_lower for k in ["guide", "tip", "help"]):
            post_type = "攻略教程"
        elif any(k in flair_lower for k in ["discussion", "talk"]):
            post_type = "讨论交流"
        elif any(k in flair_lower for k in ["question", "help"]):
            post_type = "问题求助"
        elif any(k in flair_lower for k in ["art", "video", "clip"]):
            post_type = "内容分享"
    
    # 2. Detect main topics
    detected_topics = []
    for topic, keywords in TOPIC_KEYWORDS.items():
        matches = [kw for kw in keywords if kw in combined]
        if matches:
            detected_topics.append((topic, matches[:3]))
    
    # 3. Analyze comments for discussion themes
    comment_themes = []
    if comments:
        comment_text = " ".join([c.get("body", "") for c in comments[:10]]).lower()
        for topic, keywords in TOPIC_KEYWORDS.items():
            matches = [kw for kw in keywords if kw in comment_text]
            if matches and topic not in [t[0] for t in detected_topics]:
                detected_topics.append((topic, matches[:2]))
                comment_themes.append(topic)
    
    # 4. Determine sentiment/engagement
    engagement = "低"
    if score > 100 or num_comments > 50:
        engagement = "高"
    elif score > 30 or num_comments > 20:
        engagement = "中"
    
    # 5. Extract key phrases from title
    title_cn_map = {
        "Palworld": "帕鲁",
        "CS2": "CS2",
        "Valorant": "瓦罗兰特",
        "LOL": "英雄联盟",
        "DeltaForce": "三角洲",
        "TFT": "金铲铲",
    }
    game_cn = title_cn_map.get(game, game)
    
    # 6. Build detailed summary
    summary_parts = []
    
    # Post type and engagement
    summary_parts.append(f"【帖子类型】{post_type}，热度:{engagement}(↑{score}/💬{num_comments})")
    
    # Main topics
    if detected_topics:
        topics_str = "、".join([t[0] for t in detected_topics[:4]])
        summary_parts.append(f"【讨论主题】{topics_str}")
        
        # Add matched keywords as context
        for topic, matched_kws in detected_topics[:3]:
            summary_parts.append(f"  · {topic}: 涉及「{'、'.join(matched_kws)}」")
    else:
        summary_parts.append("【讨论主题】综合话题")
    
    # Comment themes if different from post topics
    if comment_themes:
        unique_themes = [t for t in comment_themes if t not in [d[0] for d in detected_topics[:3]]]
        if unique_themes:
            summary_parts.append(f"【评论区延伸】{'、'.join(unique_themes[:3])}")
    
    # Title summary in Chinese context
    summary_parts.append(f"【主贴概述】{title[:100]}")
    
    # Selftext summary
    if selftext and len(selftext) > 20:
        # Extract first meaningful sentence
        sentences = selftext.replace("\n", " ").split(". ")
        first_sentence = sentences[0][:150] if sentences else selftext[:150]
        summary_parts.append(f"【正文要点】{first_sentence}")
    
    # Comment summary
    if comments:
        top_comment = comments[0]
        summary_parts.append(f"【热评摘要】[{top_comment.get('author','')}] {top_comment.get('body','')[:100]}")
        if len(comments) > 1:
            summary_parts.append(f"【次评摘要】[{comments[1].get('author','')}] {comments[1].get('body','')[:80]}")
    
    return "\n".join(summary_parts)


def classify_post(post, game):
    """Classify a post as A (tool opportunity) or B (pain point) and add notes."""
    title = post.get("title", "").lower()
    selftext = post.get("selftext", "").lower()
    combined = f"{title} {selftext}"
    
    tags = []
    notes = []
    
    # A类: 工具引流机会
    tool_config = TOOL_OPPORTUNITY_KEYWORDS.get(game, {})
    tool_label = tool_config.get("label", "")
    tool_keywords = tool_config.get("keywords", [])
    
    matched_tool_keywords = [kw for kw in tool_keywords if kw in combined]
    if matched_tool_keywords:
        tags.append("A-工具引流机会")
        notes.append(f"[{tool_label}] 匹配关键词: {', '.join(matched_tool_keywords[:5])}")
    
    # B类: 用户痛点
    pain_config = PAIN_POINT_KEYWORDS["all"]
    
    # B1: 抱怨功能不好用
    matched_complaints = [kw for kw in pain_config["complaint"] if kw in combined]
    if matched_complaints:
        tags.append("B-抱怨功能")
        notes.append(f"用户抱怨: {', '.join(matched_complaints[:5])}")
    
    # B2: 询问功能是否存在/需求
    matched_requests = [kw for kw in pain_config["feature_request"] if kw in combined]
    if matched_requests:
        tags.append("B-功能需求")
        notes.append(f"用户需求: {', '.join(matched_requests[:5])}")
    
    # B3: 数据错误反馈
    matched_data_errors = [kw for kw in pain_config["data_error"] if kw in combined]
    if matched_data_errors:
        tags.append("B-数据错误")
        notes.append(f"数据问题: {', '.join(matched_data_errors[:5])}")
    
    # B4: 无人解决的痛点
    matched_unmet = [kw for kw in pain_config["unmet_need"] if kw in combined]
    if matched_unmet:
        tags.append("B-无人解决")
        notes.append(f"痛点: {', '.join(matched_unmet[:5])}")
    
    if not tags:
        tags.append("其他")
    
    post["classification"] = " | ".join(tags)
    post["classification_notes"] = "\n".join(notes) if notes else "无特殊标记"
    post["is_tool_opportunity"] = any(t.startswith("A") for t in tags)
    post["is_pain_point"] = any(t.startswith("B") for t in tags)
    
    return post


def fetch_rss(subreddit, sort="hot", limit=POSTS_PER_SUB):
    """Fetch posts from Reddit RSS feed (less likely to be blocked)."""
    url = f"https://www.reddit.com/r/{subreddit}/{sort}/.rss?limit={limit}"
    print(f"Fetching RSS: {url}")
    status, content = fetch_url(url, retries=3, delay=5)
    if not content:
        print(f"  RSS fetch failed for r/{subreddit} ({sort})")
        return []

    try:
        root = ET.fromstring(content)
    except ET.ParseError as e:
        print(f"  RSS parse error: {e}")
        return []

    # Reddit RSS uses Atom feed format
    ns = {
        "atom": "http://www.w3.org/2005/Atom",
        "media": "http://search.yahoo.com/mrss/",
    }

    entries = root.findall("atom:entry", ns)
    if not entries:
        # Try without namespace
        entries = root.findall("{http://www.w3.org/2005/Atom}entry")
    
    if not entries:
        # Try plain RSS format
        entries = root.findall(".//item")
        ns = {}

    print(f"  Found {len(entries)} entries in RSS")
    posts = []

    for entry in entries[:limit]:
        post = {"subreddit": subreddit}

        # Atom format
        title = entry.find("atom:title", ns)
        if title is None:
            title = entry.find("{http://www.w3.org/2005/Atom}title")
        post["title"] = title.text if title is not None else ""

        # Link
        link = entry.find("atom:link", ns)
        if link is None:
            link = entry.find("{http://www.w3.org/2005/Atom}link")
        post["url"] = link.get("href", "") if link is not None else ""

        # Author
        author = entry.find("atom:author/atom:name", ns)
        if author is None:
            author = entry.find("{http://www.w3.org/2005/Atom}author/{http://www.w3.org/2005/Atom}name")
        post["author"] = author.text if author is not None else ""

        # Content/summary
        content_elem = entry.find("atom:content", ns)
        if content_elem is None:
            content_elem = entry.find("{http://www.w3.org/2005/Atom}content")
        summary = entry.find("atom:summary", ns)
        if summary is None:
            summary = entry.find("{http://www.w3.org/2005/Atom}summary")
        
        text = ""
        if content_elem is not None and content_elem.text:
            text = content_elem.text
        elif summary is not None and summary.text:
            text = summary.text
        
        # Clean HTML tags
        text = re.sub(r'<[^>]+>', '', text)[:500]
        post["selftext"] = text

        # Published date
        published = entry.find("atom:published", ns)
        if published is None:
            published = entry.find("{http://www.w3.org/2005/Atom}published")
        post["created_utc"] = 0
        if published is not None:
            try:
                dt = datetime.datetime.fromisoformat(published.text.replace("Z", "+00:00"))
                post["created_utc"] = dt.timestamp()
            except Exception:
                pass

        # Category/flair
        category = entry.find("atom:category", ns)
        if category is None:
            category = entry.find("{http://www.w3.org/2005/Atom}category")
        post["link_flair_text"] = category.get("term", "") if category is not None else ""

        # Try to extract score and comments from the RSS feed
        post["score"] = 0
        post["num_comments"] = 0
        post["upvote_ratio"] = 0

        # Method 1: Parse from thr:count (Atom thread extension - Reddit includes comment count)
        thr_ns = "http://purl.org/syndication/thread/1.0"
        thr_total = entry.find(f"{{{thr_ns}}}total")
        if thr_total is not None and thr_total.text:
            try:
                post["num_comments"] = int(thr_total.text)
            except:
                pass

        # Method 2: Try to find score/comments in the content HTML (before stripping tags)
        raw_content = ""
        if content_elem is not None and content_elem.text:
            raw_content = content_elem.text
        elif summary is not None and summary.text:
            raw_content = summary.text

        # Reddit RSS includes HTML like: <!-- SC_OFF --> ... <!-- SC_ON -->
        # Sometimes includes score data in data attributes or links
        score_match = re.search(r'(\d+)\s*points?', raw_content, re.IGNORECASE)
        if score_match:
            post["score"] = int(score_match.group(1))
        comment_match = re.search(r'(\d+)\s*comments?', raw_content, re.IGNORECASE)
        if comment_match and post["num_comments"] == 0:
            post["num_comments"] = int(comment_match.group(1))

        # Method 3: Parse the comments URL (Reddit RSS links include comment count in URL)
        # The link href often is: https://www.reddit.com/r/sub/comments/ID/title/
        # The comments page URL can give us the count
        all_links = entry.findall("atom:link", ns) if ns else []
        for lnk in all_links:
            if lnk is None:
                continue
            # Reddit provides a "replies" link with count
            rel = lnk.get("rel", "")
            href = lnk.get("href", "")
            if "replies" in rel or "comments" in href:
                thr_count = lnk.get("{http://purl.org/syndication/thread/1.0}count")
                if thr_count:
                    try:
                        post["num_comments"] = int(thr_count)
                    except:
                        pass

        posts.append(post)

    return posts


def get_hot_posts(subreddit_names, limit=POSTS_PER_SUB, game=None):
    """Get hot + new posts. Try JSON API FIRST (has scores), fallback to RSS."""
    if isinstance(subreddit_names, str):
        subreddit_names = [subreddit_names]

    for subreddit in subreddit_names:
        all_posts = []
        seen_urls = set()

        # Fetch both "hot" and "new" sorting
        for sort in ["hot", "new"]:
            # STRATEGY: Try JSON API FIRST (it has scores/comments built-in)
            json_success = False
            for json_host in ["https://www.reddit.com", "https://old.reddit.com"]:
                json_url = f"{json_host}/r/{subreddit}/{sort}.json?limit={limit}"
                print(f"  Trying JSON API: {json_url[:70]}")
                json_data = fetch_json(json_url, retries=2, delay=3)
                if json_data and "data" in json_data and "children" in json_data["data"]:
                    print(f"    JSON API SUCCESS! Got {len(json_data['data']['children'])} posts with scores")
                    for child in json_data["data"]["children"]:
                        d = child["data"]
                        post_url = f"https://www.reddit.com{d.get('permalink', '')}"
                        if post_url not in seen_urls:
                            all_posts.append({
                                "title": d.get("title", ""),
                                "author": d.get("author", ""),
                                "score": d.get("score", 0),
                                "num_comments": d.get("num_comments", 0),
                                "url": post_url,
                                "created_utc": d.get("created_utc", 0),
                                "selftext": (d.get("selftext", "") or "")[:500],
                                "link_flair_text": d.get("link_flair_text", ""),
                                "upvote_ratio": d.get("upvote_ratio", 0),
                                "subreddit": d.get("subreddit", subreddit),
                                "sort_type": sort,
                            })
                            seen_urls.add(post_url)
                    json_success = True
                    time.sleep(1)
                    break
                else:
                    print(f"    JSON API failed on {json_host}")
                    time.sleep(2)
            
            # Fallback to RSS only if JSON API failed
            if not json_success:
                print(f"  JSON API failed for r/{subreddit} ({sort}), falling back to RSS...")
                posts = fetch_rss(subreddit, sort=sort, limit=limit)
                if posts:
                    # RSS doesn't have scores - try to scrape from web page
                    print(f"    RSS got {len(posts)} posts, trying Playwright for scores...")
                    try:
                        from playwright.sync_api import sync_playwright
                        with sync_playwright() as p:
                            browser = p.chromium.launch(headless=True, args=["--no-sandbox", "--disable-gpu"])
                            page = browser.new_context(
                                user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/131.0.0.0 Safari/537.36"
                            ).new_page()
                            
                            # Visit subreddit page and scrape scores
                            web_url = f"https://www.reddit.com/r/{subreddit}/{sort}/"
                            page.goto(web_url, timeout=15000, wait_until="domcontentloaded")
                            time.sleep(3)
                            
                            # Parse post scores from the page
                            article_els = page.query_selector_all('article, shreddit-post, [data-testid="post-container"]')
                            for i, el in enumerate(article_els[:len(posts)]):
                                try:
                                    # Try to find score
                                    score_el = el.query_selector('[data-score]') or el.query_selector('.score')
                                    if score_el:
                                        score_text = score_el.get_attribute('data-score') or score_el.inner_text()
                                        score_match = re.search(r'(\d+)', score_text)
                                        if score_match and i < len(posts):
                                            posts[i]["score"] = int(score_match.group(1))
                                    
                                    # Try to find comment count
                                    comment_el = el.query_selector('[data-num-comments]') or el.query_selector('a[href*="comments"]')
                                    if comment_el:
                                        comment_text = comment_el.get_attribute('data-num-comments') or comment_el.inner_text()
                                        comment_match = re.search(r'(\d+)', comment_text)
                                        if comment_match and i < len(posts):
                                            posts[i]["num_comments"] = int(comment_match.group(1))
                                except:
                                    pass
                            
                            browser.close()
                            enriched = sum(1 for p in posts if p["score"] > 0)
                            print(f"    Playwright enriched {enriched}/{len(posts)} posts with scores")
                    except ImportError:
                        print(f"    Playwright not available, scores will be 0")
                    except Exception as e:
                        print(f"    Playwright error: {str(e)[:80]}")
                    
                    # Add sort tag and deduplicate
                    for p in posts:
                        p["sort_type"] = sort
                        url = p.get("url", "")
                        if url not in seen_urls:
                            all_posts.append(p)
                            seen_urls.add(url)
                time.sleep(2)

        if all_posts:
            # Classify each post
            if game:
                for p in all_posts:
                    classify_post(p, game)
            return all_posts, subreddit

        print(f"  r/{subreddit} completely failed, trying next fallback...")
        time.sleep(3)

    return [], subreddit_names[0] if subreddit_names else "unknown"


def get_top_comments(permalink, limit=COMMENTS_PER_POST):
    """Get top comments from a post via JSON API."""
    url = f"https://www.reddit.com{permalink}.json?limit={limit}&sort=top"
    data = fetch_json(url, retries=2, delay=3)
    if not data or len(data) < 2:
        # Try old.reddit.com
        url2 = f"https://old.reddit.com{permalink}.json?limit={limit}&sort=top"
        data = fetch_json(url2, retries=2, delay=3)
        if not data or len(data) < 2:
            return []

    comments = []
    try:
        for child in data[1]["data"]["children"]:
            c = child["data"]
            if c.get("body") in ("[deleted]", "[removed]", None):
                continue
            comments.append({
                "author": c.get("author", ""),
                "body": c.get("body", "")[:300],
                "score": c.get("score", 0),
            })
    except (KeyError, IndexError):
        pass
    return comments


def summarize_discussion(posts_with_comments):
    """Generate a text summary of main discussion topics."""
    if not posts_with_comments:
        return "No data available."

    topics = {}
    for p in posts_with_comments:
        flair = p.get("link_flair_text") or "General"
        if flair not in topics:
            topics[flair] = []
        topics[flair].append(p["title"])

    lines = []
    for flair, titles in topics.items():
        lines.append(f"[{flair}] ({len(titles)} posts)")
        for t in titles[:3]:
            lines.append(f"  - {t}")
        if len(titles) > 3:
            lines.append(f"  ... and {len(titles)-3} more")
        lines.append("")

    return "\n".join(lines)


def create_excel_report(all_data, output_dir=".", period_label=""):
    """Create Excel report matching the template format."""
    now = datetime.datetime.now(datetime.timezone(datetime.timedelta(hours=8)))
    date_str = now.strftime("%Y%m%d")
    time_str = now.strftime("%H%M")
    filename = f"reddit_monitor_{date_str}_{time_str}.xlsx"
    filepath = os.path.join(output_dir, filename)

    wb = openpyxl.Workbook()

    # Styling
    header_font = Font(bold=True, size=11, color="FFFFFF")
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    title_font = Font(bold=True, size=14)
    section_font = Font(bold=True, size=12, color="4472C4")
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin")
    )
    wrap_align = Alignment(wrap_text=True, vertical="top")

    # === Sheet 1: 汇总日报 ===
    ws_summary = wb.active
    ws_summary.title = "汇总日报"

    ws_summary["A1"] = f"Reddit社区巡查日报 {now.strftime('%Y年%m月%d日 %H:%M')} ({period_label})"
    ws_summary["A1"].font = title_font
    ws_summary.merge_cells("A1:G1")

    ws_summary["A3"] = "一、巡查概况"
    ws_summary["A3"].font = section_font

    ws_summary["A4"] = "社区"
    ws_summary["B4"] = "巡查帖子数"
    ws_summary["C4"] = "总评论数"
    ws_summary["D4"] = "工具引流机会"
    ws_summary["E4"] = "用户痛点帖"
    ws_summary["F4"] = "主要讨论话题"
    for col in ["A4", "B4", "C4", "D4", "E4", "F4"]:
        ws_summary[col].font = header_font
        ws_summary[col].fill = header_fill

    row = 5
    total_posts = 0
    total_comments = 0
    total_tool_opp = 0
    total_pain = 0
    for game, data in all_data.items():
        posts = data.get("posts", [])
        total_comments_count = sum(p.get("num_comments", 0) for p in posts)
        total_posts += len(posts)
        total_comments += total_comments_count

        tool_opp_count = sum(1 for p in posts if p.get("is_tool_opportunity"))
        pain_count = sum(1 for p in posts if p.get("is_pain_point"))
        total_tool_opp += tool_opp_count
        total_pain += pain_count

        top_titles = [p["title"] for p in posts[:5]]
        summary = " | ".join(top_titles[:3])

        ws_summary[f"A{row}"] = game
        ws_summary[f"B{row}"] = len(posts)
        ws_summary[f"C{row}"] = total_comments_count
        ws_summary[f"D{row}"] = tool_opp_count
        ws_summary[f"E{row}"] = pain_count
        ws_summary[f"F{row}"] = summary
        ws_summary[f"F{row}"].alignment = wrap_align
        row += 1

    ws_summary[f"A{row}"] = "合计"
    ws_summary[f"B{row}"] = total_posts
    ws_summary[f"C{row}"] = total_comments
    ws_summary[f"D{row}"] = total_tool_opp
    ws_summary[f"E{row}"] = total_pain
    ws_summary[f"A{row}"].font = Font(bold=True)
    row += 2

    ws_summary[f"A{row}"] = "二、分类说明"
    ws_summary[f"A{row}"].font = section_font
    row += 1
    legend = [
        ("A-工具引流机会", "帖子内容与我们的工具相关，可回复引导用户使用工具"),
        ("B-抱怨功能", "用户在抱怨某功能不好用，收集作为改进参考"),
        ("B-功能需求", "用户在询问某功能是否存在，收集作为新功能需求"),
        ("B-数据错误", "用户反馈数据错误，需检查工具数据准确性"),
        ("B-无人解决", "用户讨论某痛点但无人解决，可介入提供方案"),
    ]
    for tag, desc in legend:
        ws_summary[f"A{row}"] = tag
        ws_summary[f"B{row}"] = desc
        ws_summary.merge_cells(f"B{row}:F{row}")
        ws_summary[f"B{row}"].alignment = wrap_align
        if tag.startswith("A"):
            ws_summary[f"A{row}"].fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
        else:
            ws_summary[f"A{row}"].fill = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
        row += 1
    row += 1

    ws_summary[f"A{row}"] = "三、各社区讨论热点"
    ws_summary[f"A{row}"].font = section_font
    row += 1

    for game, data in all_data.items():
        posts = data.get("posts", [])
        ws_summary[f"A{row}"] = f"【{game}】r/{data.get('subreddit', game)}"
        ws_summary[f"A{row}"].font = Font(bold=True)
        row += 1
        for p in posts[:5]:
            post_summary = summarize_post_content(p, game)
            summary_lines = post_summary.split('\n')
            display = f"  · {p['title']} (score:{p['score']} comments:{p['num_comments']})\n"
            for line in summary_lines[:3]:
                display += f"    {line}\n"
            ws_summary[f"A{row}"] = display.strip()
            ws_summary[f"A{row}"].alignment = wrap_align
            ws_summary.row_dimensions[row].height = 60
            row += 1
        row += 1

    ws_summary.column_dimensions["A"].width = 40
    ws_summary.column_dimensions["B"].width = 15
    ws_summary.column_dimensions["C"].width = 15
    ws_summary.column_dimensions["D"].width = 15
    ws_summary.column_dimensions["E"].width = 15
    ws_summary.column_dimensions["F"].width = 60

    # === Sheet per game ===
    for game, data in all_data.items():
        ws = wb.create_sheet(title=game[:31])
        sub_name = data.get("subreddit", game)
        ws["A1"] = f"{now.strftime('%m月%d日 %H:%M')} r/{sub_name} 数据"
        ws["A1"].font = title_font
        ws.merge_cells("A1:J1")

        ws["A3"] = f"巡查帖子总数: {len(data.get('posts', []))}"
        ws.merge_cells("A3:J3")

        headers = [
            "序号", "帖子标题", "Flair", "作者", "点赞数",
            "评论数", "Upvote Ratio", "帖子链接", "排序类型",
            "分类标记", "分类备注/引流建议", "评论摘要(Top5)", "讨论内容总结"
        ]
        header_row = 5
        for col_idx, h in enumerate(headers, 1):
            cell = ws.cell(row=header_row, column=col_idx, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = thin_border
            cell.alignment = Alignment(horizontal="center", vertical="center")

        posts = data.get("posts", [])
        for i, post in enumerate(posts, 1):
            r = header_row + i
            comments = post.get("comments", [])
            comment_summary = "\n---\n".join(
                f"[{c['author']}] (score:{c['score']}): {c['body'][:150]}"
                for c in comments[:5]
            )

            ws.cell(row=r, column=1, value=i).border = thin_border
            ws.cell(row=r, column=2, value=post["title"]).alignment = wrap_align
            ws.cell(row=r, column=3, value=post.get("link_flair_text", ""))
            ws.cell(row=r, column=4, value=post["author"])
            ws.cell(row=r, column=5, value=post["score"])
            ws.cell(row=r, column=6, value=post["num_comments"])
            ws.cell(row=r, column=7, value=post.get("upvote_ratio", 0))
            ws.cell(row=r, column=8, value=post["url"])
            ws.cell(row=r, column=9, value=post.get("sort_type", "hot")).alignment = wrap_align
            ws.cell(row=r, column=10, value=post.get("classification", "其他")).alignment = wrap_align

            # Classification notes with color coding
            notes_cell = ws.cell(row=r, column=11, value=post.get("classification_notes", "无特殊标记"))
            notes_cell.alignment = wrap_align
            if post.get("is_tool_opportunity"):
                notes_cell.fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")  # Green
            elif post.get("is_pain_point"):
                notes_cell.fill = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")  # Yellow

            ws.cell(row=r, column=12, value=comment_summary).alignment = wrap_align
            # Column 13: Detailed content summary
            content_summary = summarize_post_content(post, game)
            ws.cell(row=r, column=13, value=content_summary).alignment = wrap_align
            # Make the summary cell taller for readability
            ws.row_dimensions[r].height = 80

            for col in range(1, 14):
                ws.cell(row=r, column=col).border = thin_border

        col_widths = {
            "A": 6, "B": 50, "C": 15, "D": 15, "E": 10,
            "F": 10, "G": 12, "H": 40, "I": 10, "J": 18,
            "K": 45, "L": 60, "M": 40
        }
        for col, width in col_widths.items():
            ws.column_dimensions[col].width = width

        ws.freeze_panes = "A6"

    wb.save(filepath)
    print(f"\nReport saved: {filepath}")
    return filepath


def get_time_period():
    """Get the monitoring time period based on REPORT_TIME env var."""
    report_time = os.environ.get("REPORT_TIME", "now")
    now_bjt = datetime.datetime.now(datetime.timezone(datetime.timedelta(hours=8)))
    
    today = now_bjt.replace(hour=0, minute=0, second=0, microsecond=0)
    
    if report_time == "09:00":
        # Report from yesterday 18:00 to today 09:00
        start = today - datetime.timedelta(days=1) + datetime.timedelta(hours=18)
        end = today + datetime.timedelta(hours=9)
        period_label = "前日18:00 - 当日09:00"
    elif report_time == "14:00":
        # Report from today 09:00 to 14:00
        start = today + datetime.timedelta(hours=9)
        end = today + datetime.timedelta(hours=14)
        period_label = "当日09:00 - 14:00"
    elif report_time == "18:00":
        # Report from today 14:00 to 18:00
        start = today + datetime.timedelta(hours=14)
        end = today + datetime.timedelta(hours=18)
        period_label = "当日14:00 - 18:00"
    else:
        # Default: last 5 hours
        start = now_bjt - datetime.timedelta(hours=5)
        end = now_bjt
        period_label = f"最近5小时 ({start.strftime('%H:%M')}-{end.strftime('%H:%M')})"
    
    return start, end, period_label


def filter_posts_by_time(posts, start_time, end_time):
    """Filter posts by created_utc within the time range."""
    start_ts = start_time.timestamp()
    end_ts = end_time.timestamp()
    
    filtered = []
    for p in posts:
        ts = p.get("created_utc", 0)
        if ts == 0:
            # If no timestamp, keep it (might be from RSS without time)
            filtered.append(p)
        elif start_ts <= ts <= end_ts:
            filtered.append(p)
    
    return filtered


def send_email(filepath, period_label, now_bjt, all_data):
    """Send the Excel report via email."""
    smtp_server = os.environ.get("SMTP_SERVER", "smtp.qq.com")
    smtp_port = int(os.environ.get("SMTP_PORT", "465"))
    email_addr = os.environ.get("EMAIL_ADDR", "851372967@qq.com")
    email_pass = os.environ.get("EMAIL_PASS", "")
    to_addr = os.environ.get("TO_ADDR", "851372967@qq.com")

    if not email_pass:
        print("EMAIL_PASS not set, skipping email")
        return False

    msg = MIMEMultipart()
    msg["From"] = email_addr
    msg["To"] = to_addr
    msg["Subject"] = f"Reddit社区巡查报告 {now_bjt.strftime('%Y-%m-%d %H:%M')} ({period_label})"

    # Build summary text
    total_posts = sum(len(d.get("posts", [])) for d in all_data.values())
    body_lines = [
        f"Reddit 社区巡查报告",
        f"时间: {now_bjt.strftime('%Y-%m-%d %H:%M')} BJT",
        f"时间段: {period_label}",
        f"总帖子数: {total_posts}",
        "",
        "各社区概况:",
    ]
    for game, data in all_data.items():
        posts = data.get("posts", [])
        tool_opp = sum(1 for p in posts if p.get("is_tool_opportunity"))
        pain = sum(1 for p in posts if p.get("is_pain_point"))
        body_lines.append(f"\n【{game}】r/{data.get('subreddit', game)} - {len(posts)}帖 (引流:{tool_opp} 痛点:{pain})")
        for p in posts[:5]:
            tag = ""
            if p.get("is_tool_opportunity"):
                tag = " [A-引流]"
            elif p.get("is_pain_point"):
                tag = " [B-痛点]"
            # Add content summary
            content_summary = summarize_post_content(p, game)
            summary_lines = content_summary.split('\n')
            body_lines.append(f"\n  {'='*50}")
            body_lines.append(f"  {p['title'][:80]}{tag}")
            for line in summary_lines[:6]:
                body_lines.append(f"  {line}")

    body = "\n".join(body_lines)
    msg.attach(MIMEText(body, "plain", "utf-8"))

    # Attach Excel file
    with open(filepath, "rb") as f:
        part = MIMEBase("application", "vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        part.set_payload(f.read())
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f'attachment; filename="{os.path.basename(filepath)}"')
        msg.attach(part)

    try:
        print(f"Sending email to {to_addr}...")
        server = smtplib.SMTP_SSL(smtp_server, smtp_port, timeout=30)
        server.login(email_addr, email_pass)
        server.sendmail(email_addr, to_addr, msg.as_string())
        server.close()
        print("Email sent successfully!")
        return True
    except Exception as e:
        print(f"Email failed: {e}", file=sys.stderr)
        return False


def create_summary_word_report(all_data, output_dir=".", period_label=""):
    """Create a separate Word document with detailed post summaries."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    now_bjt = datetime.datetime.now(datetime.timezone(datetime.timedelta(hours=8)))
    filename = f"summary_{now_bjt.strftime('%Y%m%d_%H%M')}.docx"
    filepath = os.path.join(output_dir, filename)
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    def set_shading(cell, color):
        s = OxmlElement('w:shd')
        s.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(s)
    
    def set_font(run, size=10.5, color=None, bold=False):
        run.font.size = Pt(size)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        run.bold = bold
    
    def add_heading_styled(doc, text, level=1, color='1F4E79'):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            set_font(run, size={1:16, 2:13, 3:11}.get(level, 10.5), color=color, bold=True)
        return h
    
    # === Cover ===
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('游戏社区主贴内容总结报告')
    set_font(run, size=24, color='1F4E79', bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{period_label} | {now_bjt.strftime("%Y年%m月%d日 %H:%M")} BJT')
    set_font(run, size=12, color='2E75B6')
    doc.add_page_break()
    
    # === Summary table ===
    add_heading_styled(doc, '一、巡查概况', level=1)
    table = doc.add_table(rows=len(all_data)+2, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['社区', '帖子数', 'Reddit', 'Facebook', 'X(Twitter)', '热门话题']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].add_run(h).bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_shading(cell, '1F4E79')
        set_font(cell.paragraphs[0].runs[0], size=9, color='FFFFFF', bold=True)
    
    for idx, (game, data) in enumerate(all_data.items()):
        posts = data.get("posts", [])
        rd_count = sum(1 for p in posts if p.get("sort_type") in ("hot", "new"))
        fb_count = sum(1 for p in posts if p.get("sort_type") == "facebook")
        x_count = sum(1 for p in posts if p.get("sort_type") == "x")
        
        # Collect top topics
        all_topics = []
        for p in posts:
            summary = summarize_post_content(p, game)
            for line in summary.split('\n'):
                if line.startswith('【讨论主题】'):
                    topics = line.replace('【讨论主题】', '').strip()
                    all_topics.extend(topics.split('、'))
        from collections import Counter
        top_topics = Counter(all_topics).most_common(3)
        topics_str = "、".join([t[0] for t in top_topics]) if top_topics else "综合话题"
        
        row = table.rows[idx+1]
        row.cells[0].text = game
        row.cells[1].text = str(len(posts))
        row.cells[2].text = str(rd_count)
        row.cells[3].text = str(fb_count)
        row.cells[4].text = str(x_count)
        row.cells[5].text = topics_str
        if idx % 2 == 1:
            for cell in row.cells:
                set_shading(cell, 'E7F0FA')
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    set_font(run, size=9.5)
    
    # Total row
    total_row = table.rows[len(all_data)+1]
    total_row.cells[0].text = "合计"
    total_posts = sum(len(d.get("posts", [])) for d in all_data.values())
    total_row.cells[1].text = str(total_posts)
    for cell in total_row.cells:
        set_shading(cell, 'D6E4F0')
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                set_font(run, size=9.5)
    
    doc.add_page_break()
    
    # === Detailed summaries per game ===
    add_heading_styled(doc, '二、各社区主贴详细总结', level=1)
    
    subreddit_map = {
        "Palworld": "r/Palworld", "CS2": "r/GlobalOffensive",
        "Valorant": "r/Valorant", "LOL": "r/leagueoflegends",
        "DeltaForce": "r/DeltaForce", "TFT": "r/TeamfightTactics",
    }
    
    platform_colors = {"hot": "FF6B35", "new": "FF6B35", "facebook": "1877F2", "x": "000000"}
    platform_names = {"hot": "Reddit-Hot", "new": "Reddit-New", "facebook": "Facebook", "x": "X/Twitter"}
    
    for game_idx, (game, data) in enumerate(all_data.items()):
        posts = data.get("posts", [])
        add_heading_styled(doc, f'2.{game_idx+1} {game} ({subreddit_map.get(game, game)}) - {len(posts)}帖', level=2)
        
        for post_idx, post in enumerate(posts):
            # Post title with index
            platform = post.get("sort_type", "hot")
            platform_color = platform_colors.get(platform, "808080")
            platform_name = platform_names.get(platform, platform)
            
            p = doc.add_paragraph()
            run = p.add_run(f' [{platform_name}] ')
            set_font(run, size=9, color='FFFFFF', bold=True)
            pPr = p._element.get_or_add_pPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), platform_color)
            pPr.append(shading)
            run = p.add_run(f' #{post_idx+1} {post.get("title", "")[:80]}')
            set_font(run, size=11, bold=True)
            p.paragraph_format.space_before = Pt(12)
            
            # Classification tag
            cls = post.get("classification", "")
            if cls and cls != "其他":
                p = doc.add_paragraph()
                run = p.add_run(f'  分类: {cls}')
                set_font(run, size=9, color='E74C3C' if "A" in cls else 'F39C12', bold=True)
            
            # Detailed summary
            content_summary = summarize_post_content(post, game)
            for line in content_summary.split('\n'):
                p = doc.add_paragraph()
                run = p.add_run(line)
                set_font(run, size=9.5)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.left_indent = Cm(0.5)
                
                # Color code section headers
                if line.startswith('【帖子类型】'):
                    set_font(run, size=10, color='2E86C1', bold=True)
                elif line.startswith('【讨论主题】'):
                    set_font(run, size=10, color='27AE60', bold=True)
                elif line.startswith('  · '):
                    set_font(run, size=9, color='7F8C8D')
                elif line.startswith('【评论区延伸】'):
                    set_font(run, size=9.5, color='8E44AD', bold=True)
                elif line.startswith('【热评') or line.startswith('【次评'):
                    set_font(run, size=9, color='34495E')
            
            # Post link
            p = doc.add_paragraph()
            run = p.add_run(f'  链接: {post.get("url", "")}')
            set_font(run, size=8, color='3498DB')
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(8)
        
        doc.add_page_break()
    
    # === Topic statistics ===
    add_heading_styled(doc, '三、话题统计', level=1)
    
    all_topic_counts = {}
    for game, data in all_data.items():
        posts = data.get("posts", [])
        game_topics = []
        for p in posts:
            summary = summarize_post_content(p, game)
            for line in summary.split('\n'):
                if line.startswith('【讨论主题】'):
                    topics = line.replace('【讨论主题】', '').strip()
                    game_topics.extend(topics.split('、'))
        
        add_heading_styled(doc, f'{game}', level=2, color='2E75B6')
        topic_counter = Counter(game_topics)
        table = doc.add_table(rows=len(topic_counter)+1, cols=3)
        table.style = 'Table Grid'
        headers = ['话题', '帖子数', '占比']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.paragraphs[0].add_run(h).bold = True
            set_shading(cell, '2E75B6')
            set_font(cell.paragraphs[0].runs[0], size=9, color='FFFFFF', bold=True)
        
        total = sum(topic_counter.values()) or 1
        for idx, (topic, count) in enumerate(topic_counter.most_common()):
            row = table.rows[idx+1]
            row.cells[0].text = topic
            row.cells[1].text = str(count)
            row.cells[2].text = f"{count/total*100:.0f}%"
            if idx % 2 == 1:
                for cell in row.cells:
                    set_shading(cell, 'E7F0FA')
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        set_font(run, size=9)
        
        doc.add_paragraph()
    
    doc.save(filepath)
    print(f"Summary report saved: {filepath}")
    return filepath


def send_summary_email(filepath, period_label, now_bjt, all_data):
    """Send the Word summary report via email (separate from raw data email)."""
    smtp_server = os.environ.get("SMTP_SERVER", "smtp.qq.com")
    smtp_port = int(os.environ.get("SMTP_PORT", "465"))
    email_addr = os.environ.get("EMAIL_ADDR", "851372967@qq.com")
    email_pass = os.environ.get("EMAIL_PASS", "")
    to_addr = os.environ.get("TO_ADDR", "851372967@qq.com")

    if not email_pass:
        print("EMAIL_PASS not set, skipping summary email")
        return False

    msg = MIMEMultipart()
    msg["From"] = email_addr
    msg["To"] = to_addr
    msg["Subject"] = f"主贴内容总结报告 {now_bjt.strftime('%Y-%m-%d %H:%M')} ({period_label})"

    total_posts = sum(len(d.get("posts", [])) for d in all_data.values())
    body_lines = [
        f"游戏社区主贴内容总结报告",
        f"时间: {now_bjt.strftime('%Y-%m-%d %H:%M')} BJT",
        f"时间段: {period_label}",
        f"总帖子数: {total_posts}",
        "",
        "本邮件附件为 Word 格式的主贴内容总结报告，",
        "包含每条主贴的详细内容分析（帖子类型、讨论主题、评论区分析、热度等级）。",
        "原始爬取数据请查看另一封邮件中的 Excel 附件。",
        "",
        "各社区帖子数:",
    ]
    for game, data in all_data.items():
        posts = data.get("posts", [])
        body_lines.append(f"  {game}: {len(posts)}帖")

    body = "\n".join(body_lines)
    msg.attach(MIMEText(body, "plain", "utf-8"))

    with open(filepath, "rb") as f:
        part = MIMEBase("application", "vnd.openxmlformats-officedocument.wordprocessingml.document")
        part.set_payload(f.read())
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f'attachment; filename="{os.path.basename(filepath)}"')
        msg.attach(part)

    try:
        print(f"Sending summary email to {to_addr}...")
        server = smtplib.SMTP_SSL(smtp_server, smtp_port, timeout=30)
        server.login(email_addr, email_pass)
        server.sendmail(email_addr, to_addr, msg.as_string())
        server.close()
        print("Summary email sent successfully!")
        return True
    except Exception as e:
        print(f"Summary email failed: {e}", file=sys.stderr)
        return False


def main():
    print(f"=== Reddit Community Monitor ===")
    now_bjt = datetime.datetime.now(datetime.timezone(datetime.timedelta(hours=8)))
    print(f"Time: {now_bjt.strftime('%Y-%m-%d %H:%M:%S')} BJT")

    start_time, end_time, period_label = get_time_period()
    print(f"Report period: {period_label}")
    print(f"  From: {start_time.strftime('%Y-%m-%d %H:%M')}")
    print(f"  To:   {end_time.strftime('%Y-%m-%d %H:%M')}")
    print()

    all_data = {}

    for game, subreddit_names in SUBREDDITS.items():
        print(f"\n--- Fetching {game} (subreddits: {subreddit_names}) ---")
        posts, actual_sub = get_hot_posts(subreddit_names, POSTS_PER_SUB, game=game)

        # Get top comments for each post
        for i, post in enumerate(posts):
            print(f"  [{i+1}/{len(posts)}] {post['title'][:60]}...")
            permalink = post["url"].replace("https://www.reddit.com", "").replace("https://old.reddit.com", "")
            if permalink:
                comments = get_top_comments(permalink, COMMENTS_PER_POST)
                post["comments"] = comments
                time.sleep(1)
            else:
                post["comments"] = []

        # Search Facebook via DuckDuckGo (may fail if IP blocked)
        fb_query = FB_SEARCH_QUERIES.get(game, game)
        fb_posts = search_facebook_posts(game, fb_query, limit=15)

        # Search X via DuckDuckGo (may fail if IP blocked)
        x_query = X_SEARCH_QUERIES.get(game, game)
        x_posts = search_x_posts(game, x_query, limit=15)

        # NEW: Fetch X posts via syndication API (no auth, bypasses IP blocking)
        x_api_posts = fetch_x_via_syndication(game, limit=15)

        # NEW: Find FB/X content via Reddit cross-post search
        cross_posts = fetch_fb_x_via_reddit(game, fb_query, limit=10)

        # NEW: Fetch FB/X via Playwright browser (bypasses search engine blocking)
        playwright_posts = fetch_fb_x_via_playwright(game, fb_query, x_query, limit=10)

        # Merge all: Reddit + FB search + X search + X API + cross-posts + Playwright
        all_posts = posts + fb_posts + x_posts + x_api_posts + cross_posts + playwright_posts

        # Filter posts by time period
        filtered_posts = filter_posts_by_time(all_posts, start_time, end_time)
        # If time filter removes everything (Facebook posts don't have timestamps), keep all
        if len(filtered_posts) == 0 and len(all_posts) > 0:
            filtered_posts = all_posts
        print(f"  Total: {len(all_posts)} posts (R:{len(posts)} FB:{len(fb_posts)} X:{len(x_posts)}+{len(x_api_posts)}api XP:{len(cross_posts)} PW:{len(playwright_posts)}), filtered: {len(filtered_posts)}")

        all_data[game] = {
            "subreddit": actual_sub,
            "posts": filtered_posts,
            "summary": summarize_discussion(filtered_posts),
        }

        print(f"  Got {len(filtered_posts)} posts for {game}")
        time.sleep(2)

    # Generate Excel (raw data)
    output_dir = os.environ.get("OUTPUT_DIR", ".")
    os.makedirs(output_dir, exist_ok=True)
    filepath = create_excel_report(all_data, output_dir, period_label)

    # Generate Word summary report (separate document)
    try:
        summary_path = create_summary_word_report(all_data, output_dir, period_label)
    except Exception as e:
        print(f"Word summary failed: {e}")
        summary_path = None

    # Also save raw JSON
    json_path = os.path.join(output_dir, f"reddit_raw_{now_bjt.strftime('%Y%m%d_%H%M')}.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(all_data, f, ensure_ascii=False, indent=2)
    print(f"Raw data saved: {json_path}")

    # Send email with Excel (raw data)
    send_email(filepath, period_label, now_bjt, all_data)

    # Send email with Word summary (separate)
    if summary_path:
        send_summary_email(summary_path, period_label, now_bjt, all_data)

    print("\n=== Done ===")


if __name__ == "__main__":
    main()
